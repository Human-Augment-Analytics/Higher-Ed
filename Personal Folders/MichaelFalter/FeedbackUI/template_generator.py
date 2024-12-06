from docx import Document
from docx.shared import Inches
import json

class TemplateGenerator:
    def __init__(self, categories=None):
        if categories is None:
            with open("categories.json", "r") as file:
                self.categories = json.load(file)
        else:
            self.categories = categories

    def add_header_info(self, doc, reviewer_name=None, group_name=None, researcher_name=None, date=None):
        doc.add_heading('Research Feedback Template', level=1)

        # Adjust margins
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        # Add reviewer and researcher information
        doc.add_heading('Reviewer Information', level=2)
        doc.add_paragraph(f"Reviewer Name: {reviewer_name or '______________________________'}")
        doc.add_paragraph(f"Date (Semester, Week #): {date or '____________________'}")  # ToDo: Add Date to UI

        doc.add_heading('Researcher Information', level=2)
        doc.add_paragraph(f"Researcher Name: {researcher_name or '____________________________'}")
        doc.add_paragraph(f"Researcher Group: {group_name or '___________________________'}")

        # Add a description
        doc.add_paragraph(
            "This document contains a template to record feedback for research groups. "
            "Each category represents a different metric to evaluate the quality of the research reports. "
            "Please provide scores and comments based on the criteria described."
        )
        return doc

    def add_categories(self, doc):
        # Add categories and metrics
        for category, metrics in self.categories.items():
            doc.add_heading(category, level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Score (1-5)'
            hdr_cells[2].text = 'Comments'

            for metric in metrics:
                row_cells = table.add_row().cells
                row_cells[0].text = metric

            # Add a space between categories
            doc.add_paragraph()
        return doc

    def generate_template(self, output_path, reviewer_name=None, group_name=None, researcher_name=None, date=None):
        doc = Document()
        # Add header, metadata, and description
        doc = self.add_header_info(doc, reviewer_name, group_name, researcher_name, date)
        doc = self.add_categories(doc)
        doc.save(output_path)
        print(f"Template saved at {output_path}")
