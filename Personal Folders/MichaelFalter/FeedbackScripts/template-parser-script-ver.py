import os
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

def parse_feedback_doc(doc_path):
    """
    Parses a Word document that should be a template from feedback-template-script-ver.py
    and extracts reviewer, researcher, and feedback category information.

    Parameters:
        doc_path (str): The path to the Word document file (.docx).

    Returns:
        dict: A dictionary with extracted data organized into 'Reviewer Information',
              'Researcher Information', and 'Feedback Categories'.
    """

    doc = Document(doc_path)
    
    # Dictionary to store the extracted data
    feedback_data = {
        'Reviewer Information': {},
        'Researcher Information': {},
        'Feedback Categories': {}
    }
    
    current_category = None
    tables = doc.tables
    category_keys = []

    # Helper variables to detect and parse the reviewer and researcher sections
    is_reviewer_section = False
    is_researcher_section = False

    # Iterate over paragraphs to get the reviewer, researcher, and category information
    for para in doc.paragraphs:
        text = para.text.strip()

        if text == 'Reviewer Information':
            is_reviewer_section = True
            is_researcher_section = False
            continue
        elif text == 'Researcher Information':
            is_reviewer_section = False
            is_researcher_section = True
            continue
        elif text in ['Clarity and Presentation Metrics', 
                      'Content Understanding and Communication Metrics',
                      'Technical and Analytical Depth Metrics', 
                      'Progress and Responsiveness Metrics', 
                      'Engagement and Interaction Metrics']:
            # Identify category headings
            current_category = text
            feedback_data['Feedback Categories'][current_category] = []
            category_keys.append(current_category)
            continue

        # Extract reviewer and researcher information
        if is_reviewer_section:
            if 'Reviewer Name:' in text:
                feedback_data['Reviewer Information']['Reviewer Name'] = text.replace('Reviewer Name:', '').replace('_', '').strip()
            elif 'Date (Semester, Week #):' in text:
                feedback_data['Reviewer Information']['Date'] = text.replace('Date (Semester, Week #):', '').replace('_', '').strip()
        elif is_researcher_section:
            if 'Researcher Name:' in text:
                feedback_data['Researcher Information']['Researcher Name'] = text.replace('Researcher Name:', '').replace('_', '').strip()
            elif 'Researcher Group:' in text:
                feedback_data['Researcher Information']['Researcher Group'] = text.replace('Researcher Group:', '').replace('_', '').strip()

    # Get tables from the document and associate them with detected categories
    for i, table in enumerate(tables):
        if i < len(category_keys):
            current_category = category_keys[i]
            
            # Skip header row
            for row in table.rows[1:]:
                metric = row.cells[0].text.strip()
                score = row.cells[1].text.strip()
                comments = row.cells[2].text.strip()
                
                # Store data
                feedback_data['Feedback Categories'][current_category].append({
                    'Metric': metric,
                    'Score': float(score) if score.isdigit() else None,
                    'Comments': comments
                })

    return feedback_data

def collect_feedback_data():
    """
    Collects feedback data from all Word documents (.docx) in the current directory.
    Calls `parse_feedback_doc` on each document and stores the results in a list.
    It is up to the user to ensure the docs are of the right template format.

    Returns:
        list: A list of dictionaries containing feedback data for each document.
    """
    feedback_list = []
    
    # Collects templates found in given directory. Make sure only templates are in directory.
    for file in os.listdir():
        if file.endswith('.docx'):
            feedback_list.append(parse_feedback_doc(file))
    
    return feedback_list

def calculate_average_scores_by_week(feedback_data_list):
    """
    Calculates average scores for each feedback metric, categorized by week and category.

    Parameters:
        feedback_data_list (list): A list of dictionaries with feedback data, as returned
                                   by `collect_feedback_data`.

    Returns:
        dict: A dictionary where keys are 'Semester, Week' strings (e.g., 'Fall 24, Week 11'),
              and values are dictionaries with categories as keys, each containing
              dictionaries of metrics and their average scores.
    """
    scores_by_week = {}

    for feedback_data in feedback_data_list:
        date = feedback_data['Reviewer Information'].get('Date', 'Unknown')
        week = date if date else 'Unknown'

        if week not in scores_by_week:
            scores_by_week[week] = {}

        for category, metrics in feedback_data['Feedback Categories'].items():
            if category not in scores_by_week[week]:
                scores_by_week[week][category] = {}

            for metric_data in metrics:
                metric = metric_data['Metric']
                score = metric_data['Score']

                if score is not None:
                    if metric not in scores_by_week[week][category]:
                        scores_by_week[week][category][metric] = []
                    scores_by_week[week][category][metric].append(score)

    # Calculate average scores for each week
    average_scores_by_week = {}
    for week, categories in scores_by_week.items():
        average_scores_by_week[week] = {}
        for category, metrics in categories.items():
            average_scores_by_week[week][category] = {
                metric: sum(scores) / len(scores) for metric, scores in metrics.items()
            }

    return average_scores_by_week

def plot_average_scores_by_week(average_scores_by_week):
    """
    Plots bar charts of average scores for each feedback category and metric by week.
    Result is a subplot for each category by Semester and Week.

   Parameters:
       average_scores_by_week (dict): A dictionary of average scores categorized by
                                      week and feedback category, as returned by
                                      `calculate_average_scores_by_week`.
   """

    for week, categories in average_scores_by_week.items():
        num_categories = len(categories)
        fig, axes = plt.subplots(1, num_categories, figsize=(5 * num_categories, 6))
        fig.suptitle(f'Average Scores for {week}', fontsize=16)
        
        if num_categories == 1:
            axes = [axes]  # Make axes iterable if there's only one category

        for ax, (category, metrics) in zip(axes, categories.items()):
            ax.bar(metrics.keys(), metrics.values())
            ax.set_title(category)
            ax.set_xlabel('Metric')
            ax.set_ylabel('Average Score')
            ax.set_xticklabels(metrics.keys(), rotation=45, ha='right')
            ax.set_ylim(0, 5)  # Assuming scores are between 1 and 5, up to user to follow template

        plt.tight_layout(rect=[0, 0, 1, 0.95])  
        plt.show()


def plot_metrics_over_time(average_scores_by_week, semester, category):
    """
    Plots a line chart of feedback scores for each metric over time, based on a given semester and category.

    Parameters:
        average_scores_by_week (dict): A dictionary of average scores categorized by
                                       week and feedback category, as returned by
                                       `calculate_average_scores_by_week`.
        semester (str): The semester to filter data by (e.g., 'Fall 24').
        category (str): The feedback category to filter data by (e.g., 'Clarity and Presentation Metrics').

    Returns:
        None: Displays a line plot of metric scores over time.
    """

    scores_over_time = {}

    # Extract relevant data for the semester and category
    for week_key, categories in average_scores_by_week.items():
        if semester in week_key and category in categories:
            week_number = int(week_key.split('Week ')[-1])  # Extract the week number
            for metric, score in categories[category].items():
                if metric not in scores_over_time:
                    scores_over_time[metric] = []
                scores_over_time[metric].append((week_number, score))

    # Sort the data by week number for each metric
    for metric in scores_over_time:
        scores_over_time[metric] = sorted(scores_over_time[metric], key=lambda x: x[0])

    # Prepare the plot
    plt.figure(figsize=(10, 6))
    for metric, weekly_scores in scores_over_time.items():
        weeks, scores = zip(*weekly_scores)  # Unzip week numbers and scores
        plt.plot(weeks, scores, marker='o', label=metric)

    # Plot settings
    plt.title(f'Average Scores for {category} over Time ({semester})')
    plt.xlabel('Week')
    plt.ylabel('Score')
    plt.ylim(0, 5)  # Assuming scores are between 1 and 5
    plt.legend(title="Metrics")
    plt.grid(True)
    plt.show()


# Collect feedback data from all .docx files
feedback_data_list = collect_feedback_data()

if len(feedback_data_list) == 0:
    raise Exception('No feedback data found.')

# Calculate average scores by week
average_scores_by_week = calculate_average_scores_by_week(feedback_data_list)

# PLOTS

# Plot average scores by week
plot_average_scores_by_week(average_scores_by_week)

plot_metrics_over_time(average_scores_by_week, "Fall 24", "Clarity and Presentation Metrics")
