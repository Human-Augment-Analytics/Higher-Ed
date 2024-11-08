from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()
doc.add_heading('Research Feedback Template', level=1)

# Adjust margins
sections = doc.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

# Add reviewer and researcher information
doc.add_heading('Reviewer Information', level=2)
doc.add_paragraph("Reviewer Name: ______________________________")
doc.add_paragraph("Date (Semester, Week #): ____________________")

doc.add_heading('Researcher Information', level=2)
doc.add_paragraph("Researcher Name: ____________________________")
doc.add_paragraph("Researcher Group: ___________________________")

# Add a description
doc.add_paragraph(
    "This document contains a template to record feedback for research groups. "
    "Each category represents a different metric to evaluate the quality of the research reports. "
    "Please provide scores and comments based on the criteria described."
)

# Define the categories and subcategories
categories = {
    "Clarity and Presentation Metrics": [
        "Legibility of Figures",
        "Conciseness",
        "Organization and Flow",
        "Visual Appeal"
    ],
    "Content Understanding and Communication Metrics": [
        "Accessibility for Non-Experts",
        "Data Explanation",
        "Completeness of Information"
    ],
    "Technical and Analytical Depth Metrics": [
        "Accuracy of Technical Content",
        "Depth of Analysis",
        "Appropriateness of Methodology"
    ],
    "Progress and Responsiveness Metrics": [
        "Incorporation of Feedback",
        "Timeliness and Consistency",
        "Overall Improvement Over Time"
    ],
    "Engagement and Interaction Metrics": [
        "Questions Raised and Addressed",
        "Follow-up Discussions",
        "Feedback Impact"
    ]
}

# Add a table to the document for the feedback categories
for category, metrics in categories.items():
    # Add category heading
    doc.add_heading(category, level=2)
    
    # Create a table with 3 columns: Metric, Score, Comments
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Score (1-5)'
    hdr_cells[2].text = 'Comments'
    
    # Add rows for each metric
    for metric in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric

    # Add a space between categories
    doc.add_paragraph()

# Save the document
doc.save('Research_Feedback_Template.docx')

print("Research feedback template created successfully!")
